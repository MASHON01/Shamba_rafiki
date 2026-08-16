"""
DOCX loader.

Uses python-docx. Pulls both paragraph text and table content, since
crop calendars and price/cost tables in KALRO/AFA docs carry real
information that plain paragraph extraction would silently drop.
"""

from __future__ import annotations

from pathlib import Path

from app.core.exceptions import DocumentLoadError
from app.ingestion.loaders.base import BaseLoader


class DocxLoader(BaseLoader):
    """
    Loads text content from DOCX files using python-docx.
    """

    supported_extensions = (".docx",)

    def load(self, path: Path) -> str:
        try:
            import docx
        except ImportError as exc:
            raise DocumentLoadError(
                "python-docx is required to load DOCX files. "
                "Install it with: pip install python-docx"
            ) from exc

        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise DocumentLoadError(
                f"Could not open DOCX '{path.name}': {exc}"
            ) from exc

        try:
            parts: list[str] = []

            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text.strip())

            for table in document.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if cells:
                        parts.append(" | ".join(cells))

            # Deliberately not raising here on empty `parts`: it's
            # validator.validate_text()'s job (via EmptyDocumentError)
            # to decide what counts as "usable" text - uniformly across
            # every loader, not per-format.
            return "\n\n".join(parts)

        except Exception as exc:
            raise DocumentLoadError(
                f"Could not extract text from DOCX '{path.name}': {exc}"
            ) from exc